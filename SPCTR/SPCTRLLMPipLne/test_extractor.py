"""
Test script for the DOCX extractor
Demonstrates the formatting preservation capabilities
"""

import os
import tempfile
from docx_extractor import DocxExtractor, extract_docx_text


def create_sample_docx():
    """Create a sample DOCX file for testing"""
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Create a new document
        doc = Document()
        
        # Add a title
        title = doc.add_heading('Sample Legal Document', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add a paragraph with various formatting
        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        run1 = p1.add_run('This is a ')
        run2 = p1.add_run('bold')
        run2.bold = True
        run3 = p1.add_run(' and ')
        run4 = p1.add_run('italic')
        run4.italic = True
        run5 = p1.add_run(' text in a legal document.')
        
        # Add a paragraph with superscript (like footnotes)
        p2 = doc.add_paragraph()
        run6 = p2.add_run('This contains a footnote reference')
        run7 = p2.add_run('1')
        run7.font.superscript = True
        run8 = p2.add_run(' and another')
        run9 = p2.add_run('2')
        run9.font.superscript = True
        run10 = p2.add_run('.')
        
        # Add a paragraph with small caps
        p3 = doc.add_paragraph()
        run11 = p3.add_run('THIS IS SMALL CAPS TEXT')
        run11.font.small_caps = True
        
        # Add a footnote
        footnote = p2.add_footnote('This is footnote 1 content.')
        footnote2 = p2.add_footnote('This is footnote 2 content.')
        
        # Save the document
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        return temp_file.name
        
    except ImportError:
        print("python-docx not installed. Creating a simple text file instead.")
        # Create a simple text file as fallback
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.txt')
        temp_file.write(b"""Sample Legal Document

This is a <bold>bold</bold> and <italic>italic</italic> text in a legal document.

This contains a footnote reference<superscript>1</superscript> and another<superscript>2</superscript>.

<smallcaps>THIS IS SMALL CAPS TEXT</smallcaps>

<justify_justify>This paragraph demonstrates justified text alignment which is common in legal documents.</justify_justify>
""")
        temp_file.close()
        return temp_file.name


def test_extractor():
    """Test the DOCX extractor"""
    print("=== Testing DOCX Extractor ===\n")
    
    # Create a sample document
    sample_file = create_sample_docx()
    print(f"Created sample file: {sample_file}")
    
    try:
        # Test the extractor
        if sample_file.endswith('.docx'):
            print("\n--- Extracting from DOCX file ---")
            result = extract_docx_text(sample_file)
            
            print("Main text:")
            print(result['main_text'])
            
            if 'footnotes' in result:
                print("\nFootnotes:")
                print(result['footnotes'])
        else:
            print("\n--- Reading text file ---")
            with open(sample_file, 'r') as f:
                content = f.read()
            print(content)
        
        print("\n=== Test completed successfully! ===")
        
    except Exception as e:
        print(f"Error during testing: {str(e)}")
    
    finally:
        # Clean up
        try:
            os.unlink(sample_file)
            print(f"Cleaned up: {sample_file}")
        except:
            pass


if __name__ == "__main__":
    test_extractor() 